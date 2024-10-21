# Esses imports são de todas as funcionalidades de todas as linhas do programa no total
# Tenho o pip install dessas bibliotecas no meu arquivo .txt

import os  # Não requer instalação adicional, faz parte da biblioteca padrão do Python
import json
import sys
import glob
import time
import datetime  # Essa aqui também é necessária, então sem ele dá pau
import threading
import win32print
import win32api
from pathlib import Path  # Não requer instalação adicional, faz parte da biblioteca padrão do Python
from datetime import datetime  # Não requer instalação adicional, faz parte da biblioteca padrão do Python

from pydrive.auth import GoogleAuth  # Esses dois imports são referente ao código de enviar os docs para o google drive
from pydrive.drive import GoogleDrive  # pip install PyDrive

import gspread  # Pacote para manipular o conteúdo do Google Planilhas
import pandas as pd  # pip install pandas
from openpyxl import load_workbook  # pip install openpyxl

from docxtpl import DocxTemplate  # pip install docxtpl
from docx import Document  # Para manipular o documento após renderizar o template
from docx.shared import Cm  # Para definir o tamanho da imagem no documento
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Para alinhar o parágrafo com o QR code

from oauth2client.service_account import ServiceAccountCredentials  # Necessário para que o meu código se logue no drive.
from google.oauth2.service_account import Credentials  # Alternativa para autenticar usando conta de serviço.
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.errors import HttpError


from google.oauth2 import service_account
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive

from selenium.webdriver.common.by import By  # pip install selenium
from selenium.webdriver.support.wait import WebDriverWait  # pip install selenium
from selenium import webdriver  # pip install selenium
from selenium.webdriver.support.ui import Select  # pip install selenium
from webdriver_manager.chrome import ChromeDriverManager  # pip install webdriver-manager
from selenium.webdriver.chrome.service import Service  # pip install selenium
from selenium.webdriver.support import expected_conditions as EC  # pip install selenium
from selenium.webdriver.chrome.options import Options

import qrcode  # Para gerar QR codes
from PIL import Image  # Para trabalhar com imagens
from send2trash import send2trash










# Confirmar o arquivo de credenciais da conta de serviço
def get_service_account_file():
    if getattr(sys, 'frozen', False):
        return os.path.join(sys._MEIPASS, 'service_account.json')
    else:
        return 'service_account.json'




# Função de autenticação para usar a conta de serviço no Google Drive
def autenticar_google_drive_servico():
    creds = ServiceAccountCredentials.from_json_keyfile_name(get_service_account_file(),
                                                             ['https://www.googleapis.com/auth/drive'])
    drive_service = build('drive', 'v3', credentials=creds)
    return drive_service

# Função principal para executar o código
def main():
    drive_service = autenticar_google_drive_servico()  # Autentica usando a conta de serviço
    # Aqui você pode adicionar o resto do seu código que depende da autenticação no Google Drive
    # ...

if __name__ == '__main__':
    main()

# Código Zero, copiar as informações da tabela de ofício do Google Drive
print("Copiando dados da planilha google de ofícios para o arquivo Excel. Por favor, aguarde")

# Configurações de autenticação e acesso ao Google Sheets
scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
         "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]

# Autenticação para o Google Sheets usando a conta de serviço
creds = ServiceAccountCredentials.from_json_keyfile_name("service_account.json", scope)
client = gspread.authorize(creds)

# Acessar a planilha Google pelo nome e aba específica
spreadsheet = client.open("CONTROLE DE OFÍCIOS")
sheet = spreadsheet.worksheet("TO-DO")

# Obter os dados das colunas especificadas
colunas = {
    "Tipo": 1,
    "PA": 2,
    "Referência": 3,
    "Solicitante": 4,
    "Complemento": 5,
    "Solicitado": 6,
    "Responsável": 7,
    "Cargo": 8,
    "Tipo Documento": 9,
    "Documento": 10,
    "Rua": 11,
    "GIAP": 12
}

# Mapeamento das colunas do Google Sheets para as colunas do Excel
mapeamento_excel = {
    "Tipo": 1,            # Coluna "Tipo" do Google Sheets para "CI_ou_OF" no Excel
    "PA": 3,              # Coluna "PA" do Google Sheets para a 3ª coluna no Excel
    "Referência": 4,      # Coluna "Referência" do Google Sheets para a 4ª coluna no Excel
    "Solicitante": 5,     # Coluna "Solicitante" do Google Sheets para a 5ª coluna no Excel
    "Complemento": 6,     # Coluna "Complemento" do Google Sheets para a 6ª coluna no Excel
    "Solicitado": 7,      # Coluna "Solicitado" do Google Sheets para a 7ª coluna no Excel
    "Responsável": 12,    # Coluna "Responsável" do Google Sheets para a 12ª coluna no Excel
    "Cargo": 13,          # Coluna "Cargo" do Google Sheets para a 13ª coluna no Excel
    "Tipo Documento": 9,  # Coluna "Tipo Documento" do Google Sheets para a 9ª coluna no Excel
    "Documento": 10,      # Coluna "Documento" do Google Sheets para a 10ª coluna no Excel
    "Rua": 11,            # Coluna "Rua" do Google Sheets para a 11ª coluna no Excel
    "GIAP": 14            # Coluna "GIAP" do Google Sheets para a 14ª coluna no Excel
}

# Carregar o arquivo Excel existente
excel_path = "1Informações.xlsx"
workbook = load_workbook(excel_path)
worksheet = workbook["TO-DO"]

# Função para copiar os dados de uma coluna para outra
def copiar_coluna(nome_coluna_google, indice_coluna_excel):
    dados_coluna = sheet.col_values(colunas[nome_coluna_google])[1:]  # Pega os dados da coluna, sem o cabeçalho
    for i, value in enumerate(dados_coluna, start=2):  # Começa na linha 2 para evitar sobrescrever o cabeçalho
        worksheet.cell(row=i, column=indice_coluna_excel, value=value)

# Copiar os dados conforme o mapeamento
for nome_coluna_google, indice_coluna_excel in mapeamento_excel.items():
    copiar_coluna(nome_coluna_google, indice_coluna_excel)

# Salvar o arquivo Excel modificado
workbook.save(excel_path)

# Pausa para garantir que o arquivo foi completamente salvo (ajuste conforme necessário)
time.sleep(10)  # Espera por 10 segundos (você pode ajustar este tempo)
print("Dados copiados com sucesso!")



###Código 1, referente a inserir os números de CI e Oficio na planilha


# Configuração para ocultar o navegador
#Aqui é referente ao web driver, mas já vou deixar salvo aqui mesmo.
options = Options()
options.add_argument("--headless")  # Executar em modo headless
options.add_argument("--disable-gpu")  # Necessário para Windows em alguns casos
options.add_argument("--no-sandbox")  # Necessário para algumas distribuições Linux
options.add_argument("--disable-dev-shm-usage")  # Evitar problemas de memória compartilhada
#options.add_argument("--window-size=1920,1080")  # Definir tamanho de janela padrão

servico = Service(ChromeDriverManager().install())
navegador = webdriver.Chrome(service=servico, options=options)

# Função para salvar os dados em um arquivo de texto
def salvar_dados_txt(arquivo, dados):
    with open(arquivo, 'w') as f:
        for chave, valor in dados.items():
            f.write(f"{chave},{valor}\n")

# Função para carregar os dados de um arquivo de texto
def carregar_dados_txt(arquivo):
    dados = {}
    if os.path.exists(arquivo):
        with open(arquivo, 'r') as f:
            for linha in f:
                chave, valor = linha.strip().split(',')
                dados[int(chave)] = valor  # Não converte o valor para int
    return dados

# Função para obter um novo intervalo de números
def obter_intervalo(tipo):
    inicio = int(input(f"Por favor digite o início do intervalo dos números de {tipo}: "))
    fim = int(input(f"Por favor, digite o número final do intervalo dos números de {tipo}: "))
    return inicio, fim

# Função para atribuir status inicial aos números de um intervalo
def atribuir_status(inicio, fim):
    status_numeros = {}
    for numero in range(inicio, fim + 1):
        status_numeros[numero] = "utilizável"
    return status_numeros

# Função para obter o próximo número disponível
def obter_proximo_numero(status_numeros):
    for numero, status in status_numeros.items():
        if status == "utilizável":
            return numero
    return None

# Função para marcar um número como utilizado
def marcar_numero_utilizado(status_numeros, numero):
    if numero in status_numeros:
        status_numeros[numero] = "já utilizado"

# Função para contar números disponíveis
def contar_numeros_disponiveis(status_numeros):
    return len([numero for numero, status in status_numeros.items() if status == "utilizável"])

def todos_utilizados(status):
    # Função para verificar se todos os números foram utilizados
    return all(status == 'já utilizado' for status in status.values())

# Arquivos para armazenar os dados
arquivo_ci = 'numeros_de_ci.txt'
arquivo_oficio = 'numeros_de_oficio.txt'

# Carregar dados salvos, se existirem
status_ci = carregar_dados_txt(arquivo_ci)
status_oficio = carregar_dados_txt(arquivo_oficio)

# Se não houver dados salvos ou todos os números estiverem utilizados, obter novos intervalos e salvar
if not status_ci or all(status == 'já utilizado' for status in status_ci.values()):
    inicio, fim = obter_intervalo("CI")
    status_ci = atribuir_status(inicio, fim)
    salvar_dados_txt(arquivo_ci, status_ci)

if not status_oficio or all(status == 'já utilizado' for status in status_oficio.values()):
    inicio, fim = obter_intervalo("Ofício")
    status_oficio = atribuir_status(inicio, fim)
    salvar_dados_txt(arquivo_oficio, status_oficio)

# Exibir a quantidade de números disponíveis
disponiveis_ci = contar_numeros_disponiveis(status_ci)
disponiveis_oficio = contar_numeros_disponiveis(status_oficio)

# Ler o arquivo Excel preservando a formatação
excel_file = '1Informações.xlsx'
sheet_name = 'TO-DO'

# Usar openpyxl para carregar a planilha com a formatação
wb = load_workbook(excel_file)
ws = wb[sheet_name]

# Iterar sobre as linhas da planilha
for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=2):
    tipo_cell = row[0]
    numero_doc_cell = row[1]

    if numero_doc_cell.value is None:  # Verifica se a célula está vazia
        if tipo_cell.value == 'OF':
            numero = obter_proximo_numero(status_oficio)
            if numero is not None:
                numero_doc_cell.value = numero
                marcar_numero_utilizado(status_oficio, numero)
        elif tipo_cell.value == 'CI':
            numero = obter_proximo_numero(status_ci)
            if numero is not None:
                numero_doc_cell.value = numero
                marcar_numero_utilizado(status_ci, numero)

# Salvar as mudanças no arquivo Excel
wb.save(excel_file)

# Salvar as bases de dados atualizadas
salvar_dados_txt(arquivo_ci, status_ci)
salvar_dados_txt(arquivo_oficio, status_oficio)

# Informar quantos números ainda estão disponíveis para CI e Ofício
disponiveis_ci = contar_numeros_disponiveis(status_ci)
disponiveis_oficio = contar_numeros_disponiveis(status_oficio)
print(f"Números disponíveis para CI restantes: {disponiveis_ci}")
print(f"Números disponíveis para Ofício restantes: {disponiveis_oficio}")


#Código adicional

# Verificar e solicitar números para CI e Ofício, se necessário
arquivo_ci = 'numeros_de_ci.txt'
arquivo_oficio = 'numeros_de_oficio.txt'

status_ci = carregar_dados_txt(arquivo_ci)
status_oficio = carregar_dados_txt(arquivo_oficio)

if not status_ci or todos_utilizados(status_ci):
    inicio, fim = obter_intervalo("CI")
    status_ci = atribuir_status(inicio, fim)
    salvar_dados_txt(arquivo_ci, status_ci)

if not status_oficio or todos_utilizados(status_oficio):
    inicio, fim = obter_intervalo("Ofício")
    status_oficio = atribuir_status(inicio, fim)
    salvar_dados_txt(arquivo_oficio, status_oficio)

# Exibir a quantidade de números disponíveis
disponiveis_ci = len([numero for numero, status in status_ci.items() if status == "utilizável"])
disponiveis_oficio = len([numero for numero, status in status_oficio.items() if status == "utilizável"])

#Fim código adicional
print("Números de CI e Ofício inseridos na tabela.")
###Fim Código 1











###Código 2

#Daqui para baixo é o código referente aos nomes dos secretarios
print("Aguarde enquanto os nomes dos secretários estão sendo registrados na tabela... Esse procedimento tende a demorar.")

# Carregar o arquivo Excel
workbook = load_workbook('1Informações.xlsx')

# Selecionar a planilha 'ci'
sheet = workbook['TO-DO']

# Configurar o driver do Selenium em modo headless (já configurado anteriormente)
driver = webdriver.Chrome(service=servico, options=options)




# Mapear os nomes das secretarias aos URLs correspondentes
secretaria_urls = {
    "Secretaria de Administração Geral": "https://www.carapicuiba.sp.gov.br/secretaria/view/1/administracao",
    "Secretaria de Assuntos Juridicos": "https://www.carapicuiba.sp.gov.br/secretaria/view/2/assuntos-juridicos",
    "Secretaria de Assistência Social": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/3/assistencia-social",
    "Secretaria de Cultura e Turismo": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/4/cultura-e-turismo",
    "Secretaria de Desenvolvimento Urbano": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/5/desenvolvimento-urbano",
    "Secretaria de Desenvolvimento Econômico, Social e Trabalho": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/6/desenv-economico-social-e-trabalho",
    "Secretaria de Esporte e Lazer": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/7/esporte-e-lazer",
    "Secretaria de Educação": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/8/educacao",
    "Secretaria da Fazenda": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/9/fazenda",
    "Secretaria de Infraestrutura": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/10/infraestrutura",
    "Secretaria do Meio Ambiente e Sustentabilidade": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/11/meio-ambiente-e-sustentabilidade",
    "Secretaria de Projetos Especiais, Convênios e Habitação": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/12/projetos-especiais-convenios-e-habitacao",
    "Secretaria de Governo": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/13/governo",
    "Secretaria de Obras e Serviços Municipais": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/14/obras-e-servicos-municipais",
    "Secretaria de Receita e Rendas": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/15/receita-e-rendas",
    "Secretaria de Saúde": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/16/saude",
    "Secretaria de Segurança Pública e Controle Urbano": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/17/seguranca-publica-e-controle-urbano",
    "Secretaria de Transporte e Trânsito": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/18/transporte-e-transito",
    "Gabinete do Prefeito": "https://iwww.carapicuiba.sp.gov.br/secretaria/view/19/gabinete-do-prefeito",
    # Adicione outros mapeamentos aqui conforme necessário
}


# Iterar sobre cada célula na coluna "H" (Solicitado)
for row in range(2, sheet.max_row + 1):
    Solicitado = sheet.cell(row=row, column=7).value  # Coluna "G" (Solicitado) é a coluna 7

    # Verificar se a célula não está vazia e se o nome da secretaria está mapeado
    if Solicitado and Solicitado in secretaria_urls:
        # Navegar para o URL correspondente à secretaria
        driver.get(secretaria_urls[Solicitado])

        # Encontrar o elemento com a classe 'nome-responsavel'
        nome_do_secretario = driver.find_element(By.CLASS_NAME, 'nome-responsavel').text

        # Escrever o nome do secretário na célula correspondente na coluna "H" (Nome_do_secretario)
        sheet.cell(row=row, column=8).value = nome_do_secretario  # Coluna "i" é a coluna 9

# Fechar o navegador Selenium
driver.quit()

# Mensagem de conclusão
print("Os nomes dos secretários foram registrados na tabela.")

# Salvar as alterações no arquivo Excel
workbook.save('1Informações.xlsx')
###Fim Código 2














#Código 3
#Código referente a criar o nome do documento a ser usado no giap

# Caminho para o arquivo Excel
excel_path = '1Informações.xlsx'

# Lê a planilha "TO-DO" usando pandas
df = pd.read_excel(excel_path, sheet_name='TO-DO')

# Obtém o ano atual
ano_atual = datetime.today().year

# Carrega o arquivo Excel com openpyxl para manipulação das planilhas
wb = load_workbook(excel_path)
ws_todo = wb['TO-DO']

# Itera sobre cada linha do DataFrame
for index, row in df.iterrows():
    if row['CI_ou_OF'] in ['CI', 'OF']:
        # Converte Numero_doc para inteiro
        numero_doc = int(row['Numero_doc'])
        # Formata o valor a ser inserido
        mensagem = f"{row['CI_ou_OF']} {numero_doc}-SAJ-{ano_atual}"

        # Encontra a célula na coluna "num_doc_criar_giap" correspondente à linha
        ws_todo.cell(row=index + 2, column=df.columns.get_loc('num_doc_criar_giap') + 1, value=mensagem)

# Salva o arquivo Excel
wb.save(excel_path)
#Fim do código anterior






# Código para criar o texto base que será utilizado no corpo do GIAP
# Caminho para o arquivo Excel
excel_path = '1Informações.xlsx'

# Lê a planilha "TO-DO" usando pandas
df = pd.read_excel(excel_path, sheet_name='TO-DO')

# Carrega o arquivo Excel com openpyxl para manipulação das planilhas
wb = load_workbook(excel_path)
ws_todo = wb['TO-DO']

# Itera sobre cada linha do DataFrame
for index, row in df.iterrows():
    # Verifica se a próxima linha está em branco
    if pd.isna(row['num_doc_criar_giap']) and pd.isna(row['pa']) and pd.isna(row['Referência']):
        break

    # Adiciona o prefixo "ref ao PA:" antes dos valores da coluna "pa"
    pa_value = f"ref ao PA: {row['pa']}" if not pd.isna(row['pa']) else ''

    # Copia os valores das colunas especificadas e concatena com espaços
    mensagem = f"{row['num_doc_criar_giap']} {pa_value} {row['Referência']}"

    # Encontra a célula na coluna "Mensagem_criar_giap" correspondente à linha
    ws_todo.cell(row=index + 2, column=df.columns.get_loc('Mensagem_criar_giap') + 1, value=mensagem)

# Salva o arquivo Excel
wb.save(excel_path)

print("Campos nescessarios para gerar GIAP preenchidos.")

#Fim do código anterior









#Código 4 - Referente a criar os valores de giap e os registrar na tabela



associacoes = {
    "Secretaria de Administração Geral": "SECRETARIA MUNICIPAL DA ADMINISTRAÇÃO GERAL",
    "Secretaria de Assuntos Juridicos": "SEC. MUNICIPAL DE ASSUNTOS JURIDICOS",
    "Secretaria de Assistência Social": "SEC. MUNICIPAL DE ASSISTENCIA SOCIAL E CIDADANIA",
    "Secretaria de Cultura e Turismo": "SEC. MUNICIPAL DE CULTURA E TURISMO",
    "Secretaria de Desenvolvimento Urbano": "SEC. MUNICIPAL DE DESENVOLVIMENTO URBANO E HABITACAO",
    "Secretaria de Desenvolvimento Econômico, Social e Trabalho": "SECRETARIA MUNICIPAL DE DESENVOLVIMENTO ECONÔMICO E TRABALHO",
    "Secretaria de Esporte e Lazer": "SEC. MUNICIPAL DE ESPORTE E LAZER",
    "Secretaria de Educação": "SEC. MUNICIPAL DE EDUCACAO",
    "Secretaria da Fazenda": "SEC. MUNICIPAL DA FAZENDA",
    "Secretaria de Infraestrutura": "SEC. MUNICIPAL DE INFRAESTRUTURA URBANA",
    "Secretaria do Meio Ambiente e Sustentabilidade": "SEC. MUNICIPAL DO MEIO AMBIENTE E SUSTENTABILIDADE",
    "Secretaria de Projetos Especiais, Convênios e Habitação": "SEC. MUNICIPAL DE PROJETOS ESPECIAIS E CONVENIOS",
    "Secretaria de Governo": "SEC. MUNICIPAL DE GOVERNO",
    "Secretaria de Obras e Serviços Municipais": "SEC. MUNICIPAL DE OBRAS",
    "Secretaria de Receita e Rendas": "SEC. MUNICIPAL DE RECEITA E RENDAS",
    "Secretaria de Saúde": "SEC. MUNICIPAL DE SAUDE E MEDICINA PREVENTIVA",
    "Secretaria de Segurança Pública e Controle Urbano": "SEC. MUNICIPAL DE SEGURANCA E CONTROLE URBANO",
    "Secretaria de Transporte e Trânsito": "SECRETARIA MUNICIPAL DE TRANSPORTES E TRANSITO",
    #"Secretaria de Transporte e Trânsito do município de Carapicuíba": "SECRETARIA MUNICIPAL DE TRANSPORTES E TRANSITO",

    "Gabinete do Prefeito": "GABINETE DO PREFEITO"
}


associacoes_numericas = {
    "SECRETARIA MUNICIPAL DA ADMINISTRAÇÃO GERAL": 6,
    "SEC. MUNICIPAL DE ASSUNTOS JURIDICOS": 4,
    "SEC. MUNICIPAL DE ASSISTENCIA SOCIAL E CIDADANIA": 15,
    "SEC. MUNICIPAL DE CULTURA E TURISMO": 17,
    "SEC. MUNICIPAL DE DESENVOLVIMENTO URBANO E HABITACAO": 12,
    "SECRETARIA MUNICIPAL DE DESENVOLVIMENTO ECONÔMICO E TRABALHO": 99.2,
    "SEC. MUNICIPAL DE ESPORTE E LAZER": 7,
    "SEC. MUNICIPAL DE EDUCACAO": 8,
    "SEC. MUNICIPAL DA FAZENDA": 5,
    "SEC. MUNICIPAL DE INFRAESTRUTURA URBANA": 22,
    "SEC. MUNICIPAL DO MEIO AMBIENTE E SUSTENTABILIDADE": 16,
    "SEC. MUNICIPAL DE PROJETOS ESPECIAIS E CONVENIOS": 21,
    "SEC. MUNICIPAL DE GOVERNO": 3,
    "SEC. MUNICIPAL DE OBRAS": 9,
    "SEC. MUNICIPAL DE RECEITA E RENDAS": 20,
    "SEC. MUNICIPAL DE SAUDE E MEDICINA PREVENTIVA": 13,
    "SEC. MUNICIPAL DE SEGURANCA E CONTROLE URBANO": 18,
    "SECRETARIA MUNICIPAL DE TRANSPORTES E TRANSITO": 14,
    "GABINETE DO PREFEITO": 2
}

# Carregando as informações do arquivo .json, informações ref ao GIAP
#Já modifiquei o fato de se ter também os caracteres especiais da lingua PT-BR

with open("infos_giap.json", "r", encoding="utf-8") as config_file:
    config = json.load(config_file)


# Usando os valores carregados do arquivo .json
username = config['username']
password = config['password']
secretaria_inicial_giap = config['secretaria_criadora_respectivo_giap']


# Carrega o arquivo Excel
workbook = load_workbook(filename='1Informações.xlsx')
sheet = workbook['TO-DO']

# Listas para armazenar os valores das colunas "num_doc_criar_giap", "Mensagem_criar_giap" e "Solicitado"
num_doc_values = []
mensagem_values = []
solicitado_values = []

# Encontra os índices das colunas relevantes
header = [cell.value for cell in sheet[1]]
ci_ou_of_idx = header.index("CI_ou_OF") + 1
num_doc_idx = header.index("num_doc_criar_giap") + 1
mensagem_idx = header.index("Mensagem_criar_giap") + 1
solicitado_idx = header.index("Solicitado") + 1
giap_idx = header.index("giap") + 1


# Função para atualizar a planilha
def atualizar_planilha(valor_giap_copiado):
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row):
        if row[ci_ou_of_idx - 1].value == "CI" and row[giap_idx - 1].value is None:
            row[giap_idx - 1].value = valor_giap_copiado
            workbook.save('1Informações.xlsx')
            return


# Função para processar a próxima linha
def processar_proxima_linha():
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row):
        if row[ci_ou_of_idx - 1].value == "CI" and row[giap_idx - 1].value is None:
            num_doc_values.append(row[num_doc_idx - 1].value)
            mensagem_values.append(row[mensagem_idx - 1].value)
            solicitado = row[solicitado_idx - 1].value
            associado = associacoes.get(solicitado, solicitado)
            solicitado_values.append(associado)
            break

servico = Service(ChromeDriverManager().install())
navegador = webdriver.Chrome(service=servico, options=options)

print("Os GIAPS estão sendo gerados e registrados na planilha. Esse procedimento tende a demorar...")

# Loop para processar as linhas da planilha e executar o Selenium
while True:
    processar_proxima_linha()

    if not num_doc_values or not mensagem_values or not solicitado_values:
        break


    ### Início do código webdriver-selenium

    # Acessa a página
    navegador.get('https://carapicuiba.giap.com.br/apex/carapi/f?p=652:LOGIN')
    navegador.find_element(By.ID, 'P101_USERNAME').clear()
    #navegador.find_element(By.ID, 'P101_USERNAME').send_keys('wesleygm')
    navegador.find_element(By.ID, 'P101_USERNAME').send_keys(username)

    navegador.find_element(By.ID, 'P101_PASSWORD').clear()
    navegador.find_element(By.ID, 'P101_PASSWORD').send_keys(password)
    #navegador.find_element(By.ID, 'P101_PASSWORD').send_keys('c7c2b75d')
    navegador.find_element(By.ID, 'wwvFlowForm').submit()

    navegador.find_element(By.XPATH,
                           '//*[@id="report_R5001749296453489731"]/tbody/tr[2]/td/table/tbody/tr[2]/td[2]/a').click()
    navegador.find_element(By.XPATH, '/html/body/form/div[2]/div/table/tbody/tr/td/div[1]/div[3]/div[2]/img').click()
    navegador.find_element(By.XPATH, '//*[@id="R5002551580972218898"]/tbody/tr[2]/td/ol/li[1]/a').click()
    navegador.find_element(By.XPATH, '//*[@id="P52_TIPO_EXPEDIENTE"]').click()
    select = Select(navegador.find_element(By.ID, 'P52_TIPO_EXPEDIENTE'))
    select.select_by_value('I')

    valor_para_busca = solicitado_values[0]
    search_field = navegador.find_element(By.XPATH, '//*[@id="P52_DSP_RESPONSAVEL_UNID"]')
    search_field.clear()
    search_field.send_keys(valor_para_busca)
    search_field.submit()

    navegador.find_element(By.XPATH,
                           '//*[@id="report_R110077840763092798"]/tbody/tr[2]/td/table/tbody/tr[2]/td[4]/a').click()
    select_element = navegador.find_element(By.XPATH, '//*[@id="P34_EXPE_TIPO_N"]')
    select = Select(select_element)
    select.select_by_visible_text('COMUNICAÇÃO INTERNA')


    num_doc_value = num_doc_values[0]
    campo_num_doc = navegador.find_element(By.XPATH, '//*[@id="P34_EXPE_NUM_DOCUMENTO"]')
    campo_num_doc.clear()
    campo_num_doc.send_keys(num_doc_value)

    select_element = navegador.find_element(By.XPATH,
                                            '/html/body/form/div[4]/table/tbody/tr/td/div/div[2]/div/div/table/tbody/tr/td[1]/table/tbody/tr/td/div/div[2]/div/div/table[5]/tbody/tr/td/select')
    select = Select(select_element)
    #select.select_by_visible_text('SEC. MUNICIPAL DE ASSUNTOS JURÍDICOS')
    select.select_by_visible_text(secretaria_inicial_giap)
    navegador.find_element(By.XPATH, '//*[@id="P34_COD_ASSUNTO1"]').send_keys('54')

    primeira_mensagem_value = mensagem_values[0]
    campo_mensagem_ci_values = navegador.find_element(By.XPATH, '//*[@id="P34_EXPE_DES_EXPEDIENTE"]')
    campo_mensagem_ci_values.clear()
    campo_mensagem_ci_values.send_keys(primeira_mensagem_value)
    navegador.find_element(By.XPATH, '//*[@id="B5053762502612159733"]/span').click()

    navegador.find_element(By.XPATH, '//*[@id="B5124377204040525083"]/span').click()

    codigo_da_secretaria = associacoes_numericas.get(solicitado_values[0], '')
    navegador.find_element(By.XPATH, '//*[@id="P103_UNID_COD_UNIDADE1"]').send_keys(codigo_da_secretaria)
    texto_para_tramitacao = primeira_mensagem_value
    navegador.find_element(By.XPATH, '//*[@id="P103_NOVA_COTA"]').send_keys(texto_para_tramitacao)

    copiar_giap = WebDriverWait(navegador, 1).until(
        EC.presence_of_element_located((By.ID, 'P103_DSP_NUM_EXPEDIENTE'))
    )
    valor_giap_copiado = copiar_giap.text

    # Atualiza a planilha com o valor do GIAP copiado
    atualizar_planilha(valor_giap_copiado)

    # Remove o primeiro item das listas após o processamento
    num_doc_values.pop(0)
    mensagem_values.pop(0)
    solicitado_values.pop(0)

    # Clique para tramitar o expediente
    navegador.find_element(By.XPATH, '//*[@id="B4982346863421682651"]').click()

    # Retorna para a tela principal de expediente
    navegador.find_element(By.XPATH, '//*[@id="menu_app"]').click()
    navegador.find_element(By.XPATH,
                           '//*[@id="aparece_app"]/div[2]/table/tbody/tr/td/div[2]/div/div/div[1]/div[2]/img').click()


# Fechar o navegador Selenium
driver.quit()


print("Gerando as folhas de CI'S e Oficios, aguarde.")

#Código 5 - criar de fato os documentos de CI e Oficio com a tabela preenchida e QR codes imbutidos

# Função para gerar QR Code
def generate_qr_code(data, output_path):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,  # Tamanho de cada caixa do QR Code
        border=4,  # Borda do QR Code
    )
    qr.add_data(data)
    qr.make(fit=True)

    # Criação da imagem do QR Code
    img_qr = qr.make_image(fill_color="black", back_color="white")

    # Salvando a imagem final
    img_qr.save(output_path)


# Função para inserir QR Code no documento
def insert_qr_code_in_doc(docx_path, qr_code_path, size_cm=2, position='right'):
    # Abrir o documento usando Document (depois que o template foi renderizado)
    doc = Document(docx_path)

    # Abrir a imagem do QR Code
    img = Image.open(qr_code_path)

    # Adicionar o QR Code ao documento
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Inserir a imagem com o tamanho especificado (conversão de Path para string)
    run.add_picture(str(qr_code_path), width=Cm(size_cm), height=Cm(size_cm))

    # Alinhar o parágrafo (opcional, pode mudar conforme a necessidade)
    if position == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif position == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Salvar o documento novamente com o QR code inserido
    doc.save(docx_path)


# Função para extrair o conteúdo do QR code
def generate_qr_code_content(document_type, document_number):
    # Retorna a string no formato "CI 7540" ou "OF 1340"
    return f"{document_type} {document_number}"


# Caminhos para os arquivos
base_dir = Path(__file__).parent if "__file__" in locals() else Path.cwd()
word_template_ci = base_dir / "modelo_de_ci.docx"
word_template_of = base_dir / "modelo_de_oficio.docx"
excel_path = base_dir / "1Informações.xlsx"
output_dir = base_dir / "CIS e oficios"
qr_code_dir = base_dir / "QRcodesGerados"

# Cria a pasta de saída para os documentos e QR codes
output_dir.mkdir(exist_ok=True)
qr_code_dir.mkdir(exist_ok=True)

# Converte a planilha do Excel para um dataframe do pandas
df = pd.read_excel(excel_path, sheet_name="TO-DO")

# Obtém a data atual para preencher os campos de data no modelo
dd = datetime.today().day
df["dia"] = dd

mm = datetime.today().month
meses = [
    'Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho',
    'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro'
]
mm_texto = meses[mm - 1]
df["mes"] = mm_texto

ano = datetime.today().year
df["ano"] = ano

# Remove linhas com valores NaN em "Numero_doc" e converte para inteiros
df.dropna(subset=['Numero_doc'], inplace=True)
df["Numero_doc"] = df["Numero_doc"].astype(int)

# Itera sobre cada linha do dataframe e verifica o valor de "CI_ou_OF"
for record in df.to_dict(orient="records"):
    doc = None
    output_path = None
    if record['CI_ou_OF'] == 'CI':
        if word_template_ci.exists():
            doc = DocxTemplate(word_template_ci)
        else:
            print(f"Template {word_template_ci} não encontrado.")
            continue
        output_path = output_dir / f"Comunicação Interna {record['Numero_doc']}-SAJ-{ano}.docx"
    elif record['CI_ou_OF'] == 'OF':
        if word_template_of.exists():
            doc = DocxTemplate(word_template_of)
        else:
            print(f"Template {word_template_of} não encontrado.")
            continue
        output_path = output_dir / f"OFICIO {record['Numero_doc']}-SAJ-{ano}.docx"
    else:
        continue  # Ignora linhas com valores inválidos em "CI_ou_OF"

    # Renderizar o template com os dados
    doc.render(record)
    doc.save(output_path)

    # Gerar QR Code com o tipo de documento e número
    qr_code_content = generate_qr_code_content(record['CI_ou_OF'], record['Numero_doc'])
    qr_code_path = qr_code_dir / f"{qr_code_content}_qrcode.png"
    generate_qr_code(qr_code_content, qr_code_path)

    # Inserir QR Code no documento gerado
    insert_qr_code_in_doc(output_path, qr_code_path, size_cm=2, position='right')  # Posição ajustável

    print(f"Documento gerado: {output_path}")
    print(f"QR Code gerado: {qr_code_path}")



#Excluir os QR codes

# Diretório de origem
diretorio_origem = r"C:\Users\wesley\PycharmProjects\TODO\QRcodesGerados"

# Listar todos os arquivos e subdiretórios no diretório de origem
for arquivo in os.listdir(diretorio_origem):
    caminho_arquivo = os.path.join(diretorio_origem, arquivo)
    # Mover para a lixeira
    send2trash(caminho_arquivo)

print("Todos os QR codes foram movidos para a lixeira.")



#Código 6 - Criar modelo de assinaturas para colar no caderno

# Carregar o arquivo Excel
excel_file = '1Informações.xlsx'
sheet_name = 'TO-DO'

df = pd.read_excel(excel_file, sheet_name=sheet_name)

# Obter a data atual
current_day = datetime.today().day
current_month = datetime.today().month
months = [
    'Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun',
    'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez'
]
current_month_text = months[current_month - 1]
current_year = datetime.today().year

# Carregar o modelo de documento do Word
doc = DocxTemplate('modelo de assinaturas.docx')

# Preparar os dados para preencher a tabela no Word
table_data = []
for index, row in df.iterrows():
    num_doc = row['num_doc_criar_giap']
    pa = row['pa']
    data_encaminhamento = f"{current_day} {current_month_text} {current_year}"
    table_data.append({
        'Número do Documento': num_doc,
        'PA': pa,
        'Data do encaminhamento': data_encaminhamento
    })

# Adicionar os dados na tabela do Word
context = {'table': table_data}

# Renderizar o documento com os dados
doc.render(context)

# Salvar o documento preenchido
#output_file = 'assinaturas_preenchido.docx'
#Essa linha abaixo é o meu teste
output_file = r"C:\Users\wesley\PycharmProjects\TODO\CIS e oficios\assinaturas_preenchido.docx"
doc.save(output_file)

print(f"Doc referente as assinaturas de CI's e oficios preenchido e salvo como {output_file}")





# Código 7 - Imprimir arquivos de um diretório
# Escolher qual impressora usar
lista_impressoras = win32print.EnumPrinters(2)
impressora = lista_impressoras[5]

# Mostrar a lista das impressoras conectadas ao PC
#print(lista_impressoras)

win32print.SetDefaultPrinter(impressora[2])

# Mandar imprimir todos os arquivos de uma pasta
caminho = r"C:\Users\wesley\PycharmProjects\TODO\CIS e oficios"
lista_arquivos = os.listdir(caminho)

# https://docs.microsoft.com/en-us/windows/win32/api/shellapi/nf-shellapi-shellexecutea
for arquivo in lista_arquivos:
    win32api.ShellExecute(0, "print", os.path.join(caminho, arquivo), None, caminho, 0)

print("Arquivos enviados para impressão. Aguardando a conclusão da impressão...")
# Espera para garantir que todos os arquivos sejam impressos
time.sleep(60)  # Ajuste o tempo conforme necessário para garantir que todos os arquivos sejam impressos
#Fim do código 7


# Código 7.5? enviar os arquivos para o Google Drive.
print("Os arquivos do word (de CI e OF) serão enviados para o Google Drive. Esse processo tende a demorar, aguarde.")


# Função para autenticar usando a Conta de Serviço
def authenticate_service_account():
    SCOPES = ['https://www.googleapis.com/auth/drive']

    # Substitua o caminho abaixo pelo caminho do seu arquivo JSON da conta de serviço
    service_account_file = 'service_account.json'

    # Cria as credenciais da conta de serviço
    credentials = Credentials.from_service_account_file(service_account_file, scopes=SCOPES)

    # Cria o serviço da API do Google Drive
    return build('drive', 'v3', credentials=credentials)


# Função para criar ou obter uma pasta no Google Drive
def get_or_create_folder(drive_service, parent_folder_id, folder_name):
    query = f"'{parent_folder_id}' in parents and name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    folders = results.get('files', [])

    if len(folders) > 0:
        return folders[0]['id']
    else:
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder',
            'parents': [parent_folder_id]
        }
        folder = drive_service.files().create(body=file_metadata, fields='id').execute()
        return folder['id']


# Função para fazer o upload de arquivos para o Google Drive
def upload_file_to_drive(drive_service, local_file_path, folder_id):
    try:
        file_metadata = {
            'name': os.path.basename(local_file_path),
            'parents': [folder_id]
        }
        media = MediaFileUpload(local_file_path, resumable=True)
        file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        print(f"Arquivo '{local_file_path}' enviado com sucesso para o Google Drive (ID: {file.get('id')}).")
    except HttpError as error:
        print(f"Ocorreu um erro ao enviar o arquivo '{local_file_path}': {error}")


# Função de organização e upload de arquivos
def organize_and_upload_files(drive_service, local_folder, parent_folder_id_ci, parent_folder_id_of):
    uploads = []
    # Dicionário para mapear os meses em inglês para português
    month_translation = {
        'January': 'Janeiro', 'February': 'Fevereiro', 'March': 'Março', 'April': 'Abril',
        'May': 'Maio', 'June': 'Junho', 'July': 'Julho', 'August': 'Agosto',
        'September': 'Setembro', 'October': 'Outubro', 'November': 'Novembro', 'December': 'Dezembro'
    }

    for file_name in os.listdir(local_folder):
        local_file_path = os.path.join(local_folder, file_name)

        if not os.path.isfile(local_file_path):
            continue

        creation_date = datetime.fromtimestamp(os.path.getctime(local_file_path))
        year_folder = str(creation_date.year)
        month_name = creation_date.strftime('%B')
        month_folder = f"{creation_date.month} {month_translation[month_name]}"
        day_folder = creation_date.strftime("%d-%m-%Y")

        if file_name.lower().startswith('c'):
            folder_id = get_or_create_folder(drive_service, parent_folder_id_ci, year_folder)
            folder_id = get_or_create_folder(drive_service, folder_id, month_folder)
            folder_id = get_or_create_folder(drive_service, folder_id, day_folder)
        elif file_name.lower().startswith('o'):
            folder_id = get_or_create_folder(drive_service, parent_folder_id_of, year_folder)
            folder_id = get_or_create_folder(drive_service, folder_id, month_folder)
            folder_id = get_or_create_folder(drive_service, folder_id, day_folder)
        else:
            print(f"O arquivo '{file_name}' não começa com 'c' ou 'o' e por isso não foi enviado para o drive.")
            continue

        uploads.append((drive_service, local_file_path, folder_id))

    # Enviar os arquivos
    for drive_service, local_file_path, folder_id in uploads:
        upload_file_to_drive(drive_service, local_file_path, folder_id)


# Inicializar o serviço do Google Drive usando a conta de serviço
drive_service = authenticate_service_account()

# Caminho da pasta local
local_folder = r"C:\Users\wesley\PycharmProjects\TODO\CIS e oficios"

# IDs das pastas no Google Drive (essas devem ser preenchidas com os IDs reais)
parent_folder_id_ci = '15sk1hfO8nJ8YVxHunGLaF6UQLUNtMfdu'
parent_folder_id_of = '1_q13dGGTktUiPg-saDq7RCGY9EorQKfs'

# Organizar e enviar os arquivos
organize_and_upload_files(drive_service, local_folder, parent_folder_id_ci, parent_folder_id_of)

print("Todos os arquivos foram enviados para o Google Drive com sucesso.")


# Fim Código 7.5



# Caminho dos diretórios
diretorio_oficios = r"C:\Users\wesley\PycharmProjects\TODO\CIS e oficios"
diretorio_qrcodes = r"C:\Users\wesley\PycharmProjects\TODO\QRcodesGerados"

# Padrões para localizar arquivos que começam com "o" ou "a"
padroes = [os.path.join(diretorio_oficios, '[oO]*.docx'), os.path.join(diretorio_oficios, '[aA]*.docx')]

# Enviar para a lixeira arquivos que começam com "o" ou "a"
for padrao in padroes:
    for arquivo in glob.glob(padrao):
        try:
            send2trash(arquivo)  # Envia o arquivo para a lixeira
            print(f'Arquivo movido para a lixeira: {arquivo}')
        except Exception as e:
            print(f'Erro ao mover {arquivo} para a lixeira: {e}')

# Mover/Enviar arquivos do diretório QRcodes para a lixeira
for arquivo in os.listdir(diretorio_qrcodes):
    caminho_arquivo = os.path.join(diretorio_qrcodes, arquivo)
    try:
        send2trash(caminho_arquivo)
        print(f'Arquivo movido para a lixeira: {caminho_arquivo}')
    except Exception as e:
        print(f'Erro ao mover {caminho_arquivo} para a lixeira: {e}')

print('Processo de movimentação de arquivos concluído.')

# Espera para garantir que o processo esteja completo
time.sleep(5)  # Ajuste o tempo conforme necessário









# Código 9 - Imprimir os arquivos que faltaram em um diretório
# Escolher qual impressora usar
win32print.SetDefaultPrinter(impressora[2])

# Mandar imprimir todos os arquivos restantes de uma pasta
lista_arquivos_restantes = os.listdir(caminho)

# https://docs.microsoft.com/en-us/windows/win32/api/shellapi/nf-shellapi-shellexecutea
for arquivo in lista_arquivos_restantes:
    win32api.ShellExecute(0, "print", os.path.join(caminho, arquivo), None, caminho, 0)

print("CI'S e Oficios realizados e imprimidos.")

# Espera para garantir que todos os arquivos sejam impressos
time.sleep(30)  # Ajuste o tempo conforme necessário para garantir que todos os arquivos sejam impressos

# Excluir (enviar para a lixeira) todos os arquivos restantes do diretório
for arquivo in os.listdir(caminho):
    try:
        send2trash(os.path.join(caminho, arquivo))
        print(f'Arquivo enviado para a lixeira: {os.path.join(caminho, arquivo)}')
    except Exception as e:
        print(f'Erro ao enviar {os.path.join(caminho, arquivo)} para a lixeira: {e}')

print("Todos os arquivos restantes foram enviados para a lixeira.")


#FIM CÓDIGO 9








#o código abaixo na teoria é para funcionar sem problemas, eu somente alterei a forma de autenticar no google planilhas
#Qualquer coisa, eu posso me guiar pelo código do arquivo "[alterado - final] código sem mostrar o navegador.py"

#Código 10 atualizar planilha google planilhas com os novos andamentos

# Configuração da autenticação e criação do serviço da API do Google Sheets
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
SERVICE_ACCOUNT_FILE = 'service_account.json'  # Atualize com o caminho correto para o arquivo JSON da conta de serviço

# Autenticação com a conta de serviço
creds = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES
)

# Criação do serviço do Google Sheets
service = build('sheets', 'v4', credentials=creds)

# ID da planilha Google
spreadsheet_id = '16_zlC5bRdyGTqFcVFvRIBCYzP-fjoPN9i64tD5DGe5c'

# Diretório base
base_dir = Path(__file__).parent if "__file__" in locals() else Path.cwd()

# Caminho do arquivo Excel
excel_path = base_dir / "1Informações.xlsx"

# Converte a planilha do Excel para um dataframe do pandas
df = pd.read_excel(excel_path, sheet_name="TO-DO")

# Remove linhas com valores NaN em "CI_ou_OF"
df.dropna(subset=['CI_ou_OF'], inplace=True)

# Itera sobre cada linha do dataframe
for record in df.to_dict(orient="records"):
    ci_ou_of = record['CI_ou_OF']

    # Caso 1: Se "CI_ou_OF" for "OF" e "giap" for "--"
    if ci_ou_of == "OF" and record['giap'] == "--":
        num_doc = record["Numero_doc"]  # Segunda coluna do Excel
        pa = record["pa"]  # Terceira coluna do Excel
        referencia = record["Referência"]  # Quarta coluna do Excel
        tipo_documento = record["Tipo_Documento"]  # Nona coluna do Excel
        documento = record["Documento"]  # Décima coluna do Excel
        solicitante = record["Solicitante"]  # Quinta coluna do Excel

        # Combina as colunas "Referência", "Tipo_Documento" e "Documento"
        referencia_completa = f'{referencia} {tipo_documento} {documento}'

        #Caso o meu programa de problema, só excluir essa etapa de num_doc_completo abaixo e do trecho "nova_linha = ["
        # Combina o número do documento com o texto pré-definido e o ano atual
        num_doc_completo = f"{int(num_doc)}/SAJ/{datetime.now().year}"

        # Data atual
        data_atual = datetime.now().strftime("%d/%m/%Y")

        # Dados a serem inseridos na nova linha
        nova_linha = [
            num_doc_completo, # Coluna "A" (Ofício)
            #num_doc,  # Coluna "A" (Ofício) #Vou testar a opção acima, preciso dessa modificação para usar Qr code
            #Nos meus docs
            pa,  # Coluna "B" (Processo Administrativo)
            referencia_completa,  # Coluna "C" (Referência)
            solicitante,  # Coluna "D" (Órgão de Destino)
            data_atual,  # Coluna "E" (Envio p/ Assinatura)
            ""  # Coluna "F" vazia (apaga o conteúdo)
        ]

        # Inserir a nova linha na planilha Google
        service.spreadsheets().values().append(
            spreadsheetId=spreadsheet_id,
            range="Atual!A:F",  # Atualize o nome da aba se necessário
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body={"values": [nova_linha]}
        ).execute()

        # Exibe a mensagem de registro
        print("Não foram encontrados registros anteriores, registrando o Ofício em uma nova linha.")

        # Timer para evitar bloqueio
        time.sleep(5)  # Pausa por 5 segundos

    # Caso 2: Se "CI_ou_OF" for "OF" e "giap" tiver um valor específico
    elif ci_ou_of == "OF" and record['giap'] != "--":
        num_doc = record['Numero_doc']  # Valor da coluna "Numero_doc" do Excel
        giap = record['giap']  # Valor da coluna "giap" do Excel
        pa = record["pa"]  # Terceira coluna do Excel
        referencia = record["Referência"]  # Quarta coluna do Excel
        tipo_documento = record["Tipo_Documento"]  # Nona coluna do Excel
        documento = record["Documento"]  # Décima coluna do Excel
        solicitante = record["Solicitante"]  # Quinta coluna do Excel

        # Combina as colunas "Referência", "Tipo_Documento" e "Documento"
        referencia_completa = f'{referencia} {tipo_documento} {documento}'

        # Data atual
        data_atual = datetime.now().strftime("%d/%m/%Y")

        # Caso o meu programa de problema, só excluir essa etapa de num_doc_completo abaixo e do trecho "nova_linha = ["
        # Combina o número do documento com o texto pré-definido e o ano atual
        num_doc_completo = f"{int(num_doc)}/SAJ/{datetime.now().year}"

        # Procurar o valor de "giap" na coluna GIAP (coluna "P") da planilha Google
        result = service.spreadsheets().values().get(spreadsheetId=spreadsheet_id, range="Atual!P:P").execute()
        giap_list = result.get('values', [])

        found_index = None

        for i, row in enumerate(giap_list):
            if row and row[0] == giap:
                found_index = i + 1  # +1 para ajustar ao índice da linha na planilha
                print(f"Linha encontrada com GIAP correspondente: {giap} na linha {found_index}")
                break

        if found_index:
            # Dados a serem inseridos na linha encontrada
            range_to_update = f"Atual!A{found_index}:F{found_index}"
            update_values = [
                num_doc_completo,
                #num_doc,  # Inserir o número do Oficio na linha encontrada, na coluna "A"
                pa,  # Coluna "B" (Processo Administrativo)
                referencia_completa,  # Coluna "C" (Referência)
                solicitante,  # Coluna "D" (Órgão de Destino)
                data_atual,  # Coluna "E" (Envio p/ Assinatura)
                ""  # Coluna "F" vazia (apaga o conteúdo)
            ]

            # Atualizar a linha encontrada na planilha Google
            service.spreadsheets().values().update(
                spreadsheetId=spreadsheet_id,
                range=range_to_update,
                valueInputOption="USER_ENTERED",
                body={"values": [update_values]}
            ).execute()

            # Exibe a mensagem de registro
            print(f"Atualizando os registros do oficio na linha {found_index} na planilha.")
        else:
            print(f"Nenhuma correspondência encontrada para GIAP {giap}.")

            # Criar um novo registro na planilha, assim como no Caso 1
            nova_linha = [
                num_doc_completo, # Coluna "A" (Ofício)
                #num_doc,  # Coluna "A" (Ofício)
                pa,  # Coluna "B" (Processo Administrativo)
                referencia_completa,  # Coluna "C" (Referência)
                solicitante,  # Coluna "D" (Órgão de Destino)
                data_atual,  # Coluna "E" (Envio p/ Assinatura)
                "",  # Coluna "F" vazia (apaga o conteúdo)
                # OBS A CONTAGEM DAS COLUNAS TÁ CERTO, É QUE TEM COLUNA MESCLADA NO DRIVE, POR ISSO FICA DESSA FORMA AQUI
                None,  # Coluna "G" vazia
                None,  # Coluna "H" vazia
                None,  # Coluna "i" vazia
                None,  # Coluna "J" vazia
                None,  # Coluna "K" vazia
                None,  # Coluna "L" vazia
                None,  # Coluna "M" vazia
                None,  # Coluna "N" vazia
                None,  # Coluna "O" vazia
                giap  # Coluna "P" (GIAP)
                # 9 colunas para copiar o campo de observação
            ]

            # Inserir a nova linha na planilha Google
            service.spreadsheets().values().append(
                spreadsheetId=spreadsheet_id,
                range="Atual!A:F",  # Atualize o nome da aba se necessário
                valueInputOption="USER_ENTERED",
                insertDataOption="INSERT_ROWS",
                body={"values": [nova_linha]}
            ).execute()

            # Exibe a mensagem de registro
            print("GIAP não encontrado, adicionando nova linha na planilha.")

            # Timer para evitar bloqueio
            time.sleep(5)  # Pausa por 5 segundos

    # Caso 3: Se "CI_ou_OF" for "CI"
    elif ci_ou_of == "CI":
        pa = record["pa"]  # Terceira coluna do Excel
        referencia = record["Referência"]  # Quarta coluna do Excel
        tipo_documento = record["Tipo_Documento"]  # Nona coluna do Excel
        documento = record["Documento"]  # Décima coluna do Excel
        solicitante = record["Solicitado"]  # Quinta coluna do Excel
        num_doc_criar_giap = record["num_doc_criar_giap"]
        giap_value = record["giap"]

        referencia_completa = f'{referencia}'

        data_atual = datetime.now().strftime("%d/%m/%Y")

        nova_linha = [
            None,  # Coluna "A" vazia
            pa,  # Coluna "B" (Processo Administrativo)
            referencia_completa,  # Coluna "C" (Referência)
            solicitante,  # Coluna "D" (Órgão de Destino)
            data_atual,  # Coluna "E" (Envio p/ Assinatura)
            "",  # Coluna "F" vazia (apaga o conteúdo)
            None,  # Coluna "G" vazia
            None,  # Coluna "H" vazia
            None,  # Coluna "I" vazia
            None,  # Coluna "J" vazia
            None,  # Coluna "K" vazia
            None,  # Coluna "L" vazia
            None,  # Coluna "M" vazia
            None,  # Coluna "N" vazia
            num_doc_criar_giap,  # Coluna "O" (Observação)
            giap_value  # Coluna "P" (GIAP)
        ]

        service.spreadsheets().values().append(
            spreadsheetId=spreadsheet_id,
            range="Atual!A:F",
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body={"values": [nova_linha]}
        ).execute()

        print("Registrando a CI em uma nova linha da planilha.")
        time.sleep(5)  # Pausa por 5 segundos

    # Caso 4: Se "CI_ou_OF" estiver vazio
    elif not ci_ou_of:
        pass




#Código 11 - excluir informações da planilha para a próxima execução


def limpar_planilha():
    def contagem_regressiva():
        for t in range(30, 0, -1):
            print(f"\rAs informações da planilha do Excel serão apagadas em {t} segundos...", end="")
            time.sleep(1)
        print("\nTempo esgotado. Limpando a planilha automaticamente para futuros usos.")
        realizar_limpeza()

    def realizar_limpeza():
        # Abre o arquivo Excel e seleciona a planilha "TO-DO"
        caminho_arquivo = "1Informações.xlsx"
        workbook = load_workbook(caminho_arquivo)
        planilha = workbook["TO-DO"]

        # Itera sobre as células da planilha, preservando o cabeçalho e a coluna "Q"
        for row in planilha.iter_rows(min_row=2):  # Começa a partir da segunda linha para preservar o cabeçalho
            for cell in row:
                if cell.column != 17:  # A coluna "Q" é a 17ª (coluna 17)
                    cell.value = None

        # Salva as modificações
        workbook.save(caminho_arquivo)
        print("As informações da planilha Excel foram limpas para a próxima execução.")

    # Cria um thread para a contagem regressiva
    thread_timer = threading.Thread(target=contagem_regressiva)
    thread_timer.start()

    #backup da linha abaixo
    #resposta = input("\nLimpar as informações da planilha? Responda pressionando uma das seguintes letras (s/n):").lower()

# Chama a função para execução
limpar_planilha()